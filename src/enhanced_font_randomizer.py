import os
import random
import glob
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document
from docx.oxml.ns import qn
import threading
from fontTools.ttLib import TTFont
import traceback
from docx.shared import Pt
from docx.oxml import OxmlElement
import math

class HandwritingSimulator:
    """手写模拟器 - 模拟真实手写的倾斜和纠正模式"""
    
    def __init__(self):
        self.current_tilt = 0  # 当前倾斜度
        self.char_count_since_correction = 0  # 自上次纠正后的字符数
        self.current_trend_duration = 0  # 当前趋势持续时间
        self.target_tilt = 0  # 目标倾斜度
        self.trend_direction = 0  # 趋势方向 (1: 向上, -1: 向下)
        
    def get_char_tilt(self, char):
        """
        为字符计算倾斜度
        返回: 倾斜角度（度数）
        """
        # 每10-20个字符开始新的倾斜趋势
        if (self.char_count_since_correction >= 
            random.randint(10, 20) or 
            self.current_trend_duration <= 0):
            
            # 开始新的倾斜趋势
            self._start_new_trend()
        
        # 计算当前倾斜度（平滑过渡到目标倾斜）
        progress = min(1.0, self.current_trend_duration / 15.0)
        current_tilt = self._ease_in_out(progress, self.current_tilt, self.target_tilt)
        
        # 更新状态
        self.char_count_since_correction += 1
        self.current_trend_duration -= 1
        
        # 轻微随机扰动（模拟手部微颤）
        micro_tremor = random.uniform(-0.2, 0.2)
        
        return current_tilt + micro_tremor
    
    def _start_new_trend(self):
        """开始新的倾斜趋势"""
        # 重置计数器
        self.char_count_since_correction = 0
        self.current_trend_duration = random.randint(8, 25)  # 趋势持续时间
        
        # 决定新的倾斜方向（70%概率改变方向，30%概率继续当前方向）
        if random.random() < 0.7 or abs(self.current_tilt) < 0.5:
            self.trend_direction = random.choice([-1, 1])
        else:
            # 继续当前方向但可能减弱
            self.trend_direction = 1 if self.current_tilt > 0 else -1
        
        # 设置目标倾斜度（轻微倾斜，最大1.5度）
        max_tilt = random.uniform(0.8, 1.5)
        self.target_tilt = self.trend_direction * max_tilt
        
        # 如果当前倾斜度与目标方向相反，先快速纠正
        if (self.current_tilt * self.target_tilt) < 0:
            # 方向相反，先快速回归基线
            correction_duration = random.randint(3, 8)
            self.current_trend_duration = correction_duration
            self.target_tilt = 0  # 先回归基线
            
        # 更新当前倾斜度为起始点
        self.current_tilt = self.current_tilt
    
    def _ease_in_out(self, t, start, end):
        """缓动函数，使倾斜变化更自然"""
        # 三次缓动函数
        t = max(0, min(1, t))
        t2 = t * t
        t3 = t2 * t
        return start + (end - start) * (-2 * t3 + 3 * t2)

class FontManager:
    """字体管理器，负责加载字体和检查字符可用性"""
    
    def __init__(self, fonts_dir="fonts"):
        self.fonts_dir = fonts_dir
        self.font_files = []
        self.font_cache = {}  # 缓存字体对象和字符集
        self.load_fonts()
    
    def load_fonts(self):
        """加载所有字体文件并预加载字符信息"""
        self.font_files = []
        self.font_cache = {}
        
        if os.path.exists(self.fonts_dir):
            for ext in ['*.ttf', '*.otf']:
                font_paths = glob.glob(os.path.join(self.fonts_dir, ext))
                for font_path in font_paths:
                    try:
                        # 尝试加载字体并提取字符集
                        font = TTFont(font_path)
                        font_name = os.path.splitext(os.path.basename(font_path))[0]
                        
                        # 获取字体支持的字符
                        chars = set()
                        for table in font['cmap'].tables:
                            chars.update(table.cmap.keys())
                        
                        self.font_cache[font_name] = {
                            'path': font_path,
                            'chars': chars,
                            'object': font
                        }
                        self.font_files.append(font_path)
                        
                        print(f"加载字体: {font_name} (包含 {len(chars)} 个字符)")
                        
                    except Exception as e:
                        print(f"加载字体 {font_path} 时出错: {e}")
        
        return len(self.font_files)
    
    def get_font_for_char(self, char):
        """为指定字符查找可用的字体"""
        if not self.font_cache:
            return None
            
        char_code = ord(char)
        
        # 随机打乱字体顺序，确保随机性
        font_names = list(self.font_cache.keys())
        random.shuffle(font_names)
        
        # 查找支持该字符的字体
        for font_name in font_names:
            if char_code in self.font_cache[font_name]['chars']:
                return font_name
        
        # 如果没有字体支持该字符，返回None
        return None
    
    def get_random_font_name(self):
        """获取随机字体名称"""
        if not self.font_cache:
            return None
        return random.choice(list(self.font_cache.keys()))
    
    def get_font_count(self):
        """获取字体数量"""
        return len(self.font_cache)
    
    def get_font_names(self):
        """获取所有字体名称"""
        return list(self.font_cache.keys())

class LineSpacingManager:
    """行间距管理器 - 实现每两行之间的随机间距"""
    
    def __init__(self):
        self.line_spacing_cache = {}  # 缓存已设置的行间距
        
    def get_random_line_spacing(self, line_index):
        """
        为指定行获取随机行间距
        返回: 行间距倍数 (0.8~1.2之间)
        """
        # 如果已经为这行设置过间距，则返回缓存值
        if line_index in self.line_spacing_cache:
            return self.line_spacing_cache[line_index]
        
        # 生成随机行间距
        spacing = random.uniform(0.8, 1.2)
        self.line_spacing_cache[line_index] = spacing
        return spacing

class FontRandomizerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("增强版字体随机替换工具")
        self.root.geometry("800x700")
        self.root.resizable(True, True)
        
        # 初始化变量
        self.input_file = ""
        self.output_file = ""
        self.fonts_dir = "fonts"
        self.is_processing = False
        
        # 手写效果设置
        self.enable_handwriting_effect = tk.BooleanVar(value=True)
        self.handwriting_strength = tk.IntVar(value=3)
        
        # 新功能设置
        self.enable_random_line_spacing = tk.BooleanVar(value=True)
        self.enable_random_char_size = tk.BooleanVar(value=True)
        self.enable_fine_line_spacing = tk.BooleanVar(value=True)  # 精细行间距控制
        self.enable_random_indent = tk.BooleanVar(value=True)  # 随机行首缩进
        
        # 力度调节
        self.char_size_strength = tk.IntVar(value=3)  # 字号随机力度
        self.line_spacing_strength = tk.IntVar(value=3)  # 行间距随机力度
        self.indent_strength = tk.IntVar(value=3)  # 缩进随机力度
        
        # 字符高度控制
        self.last_char_position = None  # 上一个字符的垂直位置
        
        # 创建带滚动条的主容器
        self.create_scrollable_mainframe()
        
        # 加载字体
        self.load_fonts()
        
    def create_scrollable_mainframe(self):
        """创建带滚动条的主容器"""
        # 创建主框架
        main_frame = ttk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 创建画布和滚动条
        canvas = tk.Canvas(main_frame)
        scrollbar = ttk.Scrollbar(main_frame, orient=tk.VERTICAL, command=canvas.yview)
        self.scrollable_frame = ttk.Frame(canvas)
        
        # 配置滚动区域
        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        # 创建窗口
        canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # 绑定鼠标滚轮事件
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        
        canvas.bind("<MouseWheel>", _on_mousewheel)
        self.scrollable_frame.bind("<MouseWheel>", _on_mousewheel)
        
        # 布局
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 创建界面组件
        self.create_widgets()
        
    def create_widgets(self):
        """创建界面组件"""
        # 主容器
        main_frame = ttk.Frame(self.scrollable_frame, padding="15")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 标题
        title_label = ttk.Label(
            main_frame, 
            text="增强版字符级字体随机替换工具", 
            font=("Arial", 16, "bold")
        )
        title_label.pack(pady=(0, 15))
        
        # 说明文本
        desc_text = (
            "本工具会将Word文档中的每个字符随机替换为不同字体\n"
            "新增功能：智能手写模拟效果，模拟真实手写的倾斜和纠正模式"
        )
        desc_label = ttk.Label(main_frame, text=desc_text, justify=tk.CENTER)
        desc_label.pack(pady=(0, 15))
        
        # 文件选择区域
        file_frame = ttk.LabelFrame(main_frame, text="文件选择", padding="10")
        file_frame.pack(fill=tk.X, pady=10)
        
        # 输入文件
        ttk.Label(file_frame, text="输入文件:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.input_entry = ttk.Entry(file_frame, width=60)
        self.input_entry.grid(row=0, column=1, padx=5, pady=5, sticky=tk.W+tk.E)
        ttk.Button(file_frame, text="浏览", command=self.browse_input).grid(row=0, column=2, padx=5, pady=5)
        
        # 输出文件
        ttk.Label(file_frame, text="输出文件:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.output_entry = ttk.Entry(file_frame, width=60)
        self.output_entry.grid(row=1, column=1, padx=5, pady=5, sticky=tk.W+tk.E)
        ttk.Button(file_frame, text="浏览", command=self.browse_output).grid(row=1, column=2, padx=5, pady=5)
        
        file_frame.columnconfigure(1, weight=1)
        
        # 手写效果设置区域
        handwriting_frame = ttk.LabelFrame(main_frame, text="手写效果设置", padding="10")
        handwriting_frame.pack(fill=tk.X, pady=10)
        
        # 启用手写效果
        handwriting_check = ttk.Checkbutton(
            handwriting_frame, 
            text="启用手写模拟效果（智能倾斜+纠正）", 
            variable=self.enable_handwriting_effect
        )
        handwriting_check.pack(anchor=tk.W)
        
        # 手写强度设置
        strength_frame = ttk.Frame(handwriting_frame)
        strength_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(strength_frame, text="手写自然度:").pack(side=tk.LEFT)
        strength_scale = ttk.Scale(
            strength_frame, 
            from_=1, 
            to=5, 
            variable=self.handwriting_strength,
            orient=tk.HORIZONTAL
        )
        strength_scale.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        self.strength_label = ttk.Label(strength_frame, text="中等")
        self.strength_label.pack(side=tk.RIGHT)
        
        # 绑定事件
        strength_scale.configure(command=self.update_handwriting_labels)
        
        # 新功能设置区域
        new_features_frame = ttk.LabelFrame(main_frame, text="随机化效果设置", padding="10")
        new_features_frame.pack(fill=tk.X, pady=10)
        
        # 随机行间距
        line_spacing_check = ttk.Checkbutton(
            new_features_frame, 
            text="启用随机行间距", 
            variable=self.enable_random_line_spacing
        )
        line_spacing_check.pack(anchor=tk.W, pady=2)
        
        # 行间距力度设置
        line_spacing_strength_frame = ttk.Frame(new_features_frame)
        line_spacing_strength_frame.pack(fill=tk.X, pady=2, padx=20)
        
        ttk.Label(line_spacing_strength_frame, text="行间距随机力度:").pack(side=tk.LEFT)
        line_spacing_strength_scale = ttk.Scale(
            line_spacing_strength_frame, 
            from_=1, 
            to=5, 
            variable=self.line_spacing_strength,
            orient=tk.HORIZONTAL
        )
        line_spacing_strength_scale.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        self.line_spacing_strength_label = ttk.Label(line_spacing_strength_frame, text="中等")
        self.line_spacing_strength_label.pack(side=tk.RIGHT)
        
        # 随机字符大小
        char_size_check = ttk.Checkbutton(
            new_features_frame, 
            text="启用随机字符大小", 
            variable=self.enable_random_char_size
        )
        char_size_check.pack(anchor=tk.W, pady=2)
        
        # 字符大小力度设置
        char_size_strength_frame = ttk.Frame(new_features_frame)
        char_size_strength_frame.pack(fill=tk.X, pady=2, padx=20)
        
        ttk.Label(char_size_strength_frame, text="字号随机力度:").pack(side=tk.LEFT)
        char_size_strength_scale = ttk.Scale(
            char_size_strength_frame, 
            from_=1, 
            to=5, 
            variable=self.char_size_strength,
            orient=tk.HORIZONTAL
        )
        char_size_strength_scale.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        self.char_size_strength_label = ttk.Label(char_size_strength_frame, text="中等")
        self.char_size_strength_label.pack(side=tk.RIGHT)
        
        # 随机行首缩进
        indent_check = ttk.Checkbutton(
            new_features_frame, 
            text="启用随机行首缩进（每行前加1~3个空格）", 
            variable=self.enable_random_indent
        )
        indent_check.pack(anchor=tk.W, pady=2)
        
        # 缩进力度设置
        indent_strength_frame = ttk.Frame(new_features_frame)
        indent_strength_frame.pack(fill=tk.X, pady=2, padx=20)
        
        ttk.Label(indent_strength_frame, text="缩进随机力度:").pack(side=tk.LEFT)
        indent_strength_scale = ttk.Scale(
            indent_strength_frame, 
            from_=1, 
            to=5, 
            variable=self.indent_strength,
            orient=tk.HORIZONTAL
        )
        indent_strength_scale.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        self.indent_strength_label = ttk.Label(indent_strength_frame, text="中等")
        self.indent_strength_label.pack(side=tk.RIGHT)
        
        # 绑定事件
        line_spacing_strength_scale.configure(command=self.update_strength_labels)
        char_size_strength_scale.configure(command=self.update_strength_labels)
        indent_strength_scale.configure(command=self.update_strength_labels)
        
        # 字体信息区域
        font_frame = ttk.LabelFrame(main_frame, text="字体信息", padding="10")
        font_frame.pack(fill=tk.X, pady=10)
        
        # 字体列表
        ttk.Label(font_frame, text="可用字体:").pack(anchor=tk.W)
        
        list_frame = ttk.Frame(font_frame)
        list_frame.pack(fill=tk.X, pady=5)
        
        self.font_listbox = tk.Listbox(list_frame, height=5)
        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.font_listbox.yview)
        self.font_listbox.configure(yscrollcommand=scrollbar.set)
        
        self.font_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 字体操作按钮
        font_btn_frame = ttk.Frame(font_frame)
        font_btn_frame.pack(fill=tk.X, pady=5)
        
        ttk.Button(font_btn_frame, text="刷新字体列表", command=self.load_fonts).pack(side=tk.LEFT, padx=5)
        ttk.Button(font_btn_frame, text="打开字体文件夹", command=self.open_font_folder).pack(side=tk.LEFT, padx=5)
        
        # 状态信息
        status_frame = ttk.Frame(font_frame)
        status_frame.pack(fill=tk.X, pady=5)
        
        self.font_count_label = ttk.Label(status_frame, text="检测到 0 个字体文件")
        self.font_count_label.pack(side=tk.LEFT)
        
        # 操作按钮区域
        btn_frame = ttk.Frame(main_frame)
        btn_frame.pack(fill=tk.X, pady=20)
        
        self.convert_btn = ttk.Button(
            btn_frame, 
            text="开始字符级字体替换", 
            command=self.start_conversion
        )
        self.convert_btn.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(btn_frame, text="清空", command=self.clear_all).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="退出", command=self.root.quit).pack(side=tk.RIGHT, padx=5)
        
        # 进度和日志区域
        log_frame = ttk.LabelFrame(main_frame, text="处理日志", padding="10")
        log_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        self.log_text = tk.Text(log_frame, height=12, wrap=tk.WORD)
        log_scrollbar = ttk.Scrollbar(log_frame, orient=tk.VERTICAL, command=self.log_text.yview)
        self.log_text.configure(yscrollcommand=log_scrollbar.set)
        
        self.log_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        log_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # 状态栏
        self.status_var = tk.StringVar(value="就绪")
        status_bar = ttk.Label(main_frame, textvariable=self.status_var, relief=tk.SUNKEN)
        status_bar.pack(fill=tk.X, pady=(5, 0))
        
        # 初始日志
        self.log("程序已启动")
        self.log("请选择输入和输出文件，然后点击'开始字符级字体替换'")
        
    def update_handwriting_labels(self, *args):
        """更新手写效果标签"""
        strength_texts = ["轻微", "较弱", "中等", "较强", "真实"]
        strength_value = min(max(0, self.handwriting_strength.get() - 1), 4)
        self.strength_label.config(text=strength_texts[strength_value])
        
    def update_strength_labels(self, *args):
        """更新力度标签"""
        strength_texts = ["轻微", "较弱", "中等", "较强", "强烈"]
        
        # 更新行间距力度标签
        line_spacing_value = min(max(0, self.line_spacing_strength.get() - 1), 4)
        self.line_spacing_strength_label.config(text=strength_texts[line_spacing_value])
        
        # 更新字号力度标签
        char_size_value = min(max(0, self.char_size_strength.get() - 1), 4)
        self.char_size_strength_label.config(text=strength_texts[char_size_value])
        
        # 更新缩进力度标签
        indent_value = min(max(0, self.indent_strength.get() - 1), 4)
        self.indent_strength_label.config(text=strength_texts[indent_value])
        
    def load_fonts(self):
        """加载字体文件"""
        try:
            self.log("正在加载字体文件...")
            self.update_status("正在加载字体...")
            self.root.update()
            
            self.font_manager = FontManager(self.fonts_dir)
            font_count = len(self.font_manager.font_files)
            
            # 更新字体列表
            self.font_listbox.delete(0, tk.END)
            for font_name in self.font_manager.font_cache.keys():
                self.font_listbox.insert(tk.END, font_name)
            
            # 更新状态
            self.font_count_label.config(text=f"检测到 {font_count} 个字体文件")
            
            if font_count == 0:
                self.log("警告: 没有找到字体文件！请将.ttf或.otf文件放入fonts文件夹")
            else:
                self.log(f"已加载 {font_count} 个字体文件")
                self.log("字体字符集已预加载，将确保字符可用性")
                
        except Exception as e:
            self.log(f"加载字体时出错: {str(e)}")
            self.log("请确保已安装fontTools库: pip install fonttools")
        
        self.update_status("就绪")
    
    def browse_input(self):
        """浏览输入文件"""
        filename = filedialog.askopenfilename(
            title="选择Word文档",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if filename:
            self.input_entry.delete(0, tk.END)
            self.input_entry.insert(0, filename)
            self.input_file = filename
            
            # 自动生成输出文件名
            if not self.output_entry.get():
                dir_name = os.path.dirname(filename)
                file_name = os.path.basename(filename)
                name, ext = os.path.splitext(file_name)
                output_path = os.path.join(dir_name, f"{name}_随机字体{ext}")
                self.output_entry.insert(0, output_path)
                self.output_file = output_path
            
            self.log(f"已选择输入文件: {filename}")
            self.update_status("就绪 - 可以选择输出文件并开始转换")
    
    def browse_output(self):
        """浏览输出文件"""
        filename = filedialog.asksaveasfilename(
            title="保存转换后的文档",
            defaultextension=".docx",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )
        if filename:
            self.output_entry.delete(0, tk.END)
            self.output_entry.insert(0, filename)
            self.output_file = filename
            self.log(f"已设置输出文件: {filename}")
            self.update_status("就绪 - 可以开始转换")
    
    def open_font_folder(self):
        """打开字体文件夹"""
        if not os.path.exists(self.fonts_dir):
            os.makedirs(self.fonts_dir)
            self.log("已创建fonts文件夹")
        
        try:
            os.startfile(self.fonts_dir)  # Windows
        except:
            try:
                os.system(f'open "{self.fonts_dir}"')  # Mac
            except:
                try:
                    os.system(f'xdg-open "{self.fonts_dir}"')  # Linux
                except:
                    self.log("无法打开字体文件夹")
        
        self.log("已打开字体文件夹，请将.ttf或.otf字体文件放入此文件夹")
    
    def start_conversion(self):
        """开始转换"""
        if self.is_processing:
            return
            
        input_file = self.input_entry.get().strip()
        output_file = self.output_entry.get().strip()
        
        # 验证输入
        if not input_file:
            messagebox.showerror("错误", "请选择输入文件")
            return
        
        if not output_file:
            messagebox.showerror("错误", "请设置输出文件")
            return
        
        if not os.path.exists(input_file):
            messagebox.showerror("错误", "输入文件不存在")
            return
        
        if not hasattr(self, 'font_manager') or len(self.font_manager.font_files) < 2:
            messagebox.showerror("错误", "至少需要2个字体文件才能实现字符级随机替换")
            return
        
        # 开始转换（在新线程中）
        self.is_processing = True
        self.convert_btn.config(state="disabled")
        self.update_status("正在处理...")
        
        thread = threading.Thread(
            target=self.convert_document, 
            args=(input_file, output_file)
        )
        thread.daemon = True
        thread.start()
    
    def convert_document(self, input_path, output_path):
        """转换文档（在单独线程中运行）"""
        try:
            self.log("开始字符级字体随机替换...")
            self.log("正在处理文档，请稍候...")
            
            # 统计信息
            stats = {
                'total_chars': 0,
                'chars_with_font': 0,
                'chars_without_font': 0,
                'used_fonts': set(),
                'handwriting_trends': 0,
                'lines_with_random_spacing': 0,
                'chars_with_random_size': 0,
                'lines_with_random_indent': 0
            }
            
            # 根据强度调整手写参数
            strength = self.handwriting_strength.get()
            max_tilt_multiplier = 0.5 + (strength * 0.3)  # 1.0-2.0倍倾斜
            
            # 根据力度调整随机范围
            char_size_range = 0.3 + (self.char_size_strength.get() * 0.3)  # 0.6-1.8
            line_spacing_min = 0.9 - (self.line_spacing_strength.get() * 0.1)  # 0.8-0.5
            line_spacing_max = 1.1 + (self.line_spacing_strength.get() * 0.1)  # 1.2-1.6
            
            # 根据力度调整缩进范围
            indent_strength = self.indent_strength.get()
            indent_min = 1  # 最少1个空格
            indent_max = min(5, 1 + indent_strength)  # 最多1+力度值个空格，最大5个
            
            # 加载文档
            doc = Document(input_path)
            
            # 为每个段落创建独立的手写模拟器
            paragraph_simulators = {}
            
            # 创建行间距管理器
            line_spacing_manager = LineSpacingManager()
            
            # 处理所有段落
            for paragraph_idx, paragraph in enumerate(doc.paragraphs):
                # 为每个段落创建独立的手写模拟器
                if paragraph_idx not in paragraph_simulators:
                    paragraph_simulators[paragraph_idx] = HandwritingSimulator()
                
                simulator = paragraph_simulators[paragraph_idx]
                
                # 应用随机行间距
                if self.enable_random_line_spacing.get() and paragraph.text.strip():
                    # 根据力度调整行间距范围
                    random_spacing = random.uniform(line_spacing_min, line_spacing_max)
                    paragraph.paragraph_format.line_spacing = random_spacing
                    stats['lines_with_random_spacing'] += 1
                
                # 应用随机行首缩进
                if self.enable_random_indent.get() and paragraph.text.strip():
                    # 在缩进范围内随机选择空格数量
                    indent_spaces = random.randint(indent_min, indent_max)
                    # 在段落开头添加空格
                    if paragraph.runs:
                        # 如果段落已有内容，在第一个run前插入空格
                        first_run = paragraph.runs[0]
                        spaces = " " * indent_spaces
                        first_run.text = spaces + first_run.text
                        stats['lines_with_random_indent'] += 1
                    else:
                        # 如果段落没有内容，添加一个包含空格的run
                        paragraph.add_run(" " * indent_spaces)
                        stats['lines_with_random_indent'] += 1
                
                # 初始化字符大小跟踪和高度位置跟踪
                last_char_size = None
                self.last_char_position = None
                
                runs = list(paragraph.runs)
                for run in runs:
                    text = run.text
                    if text.strip():
                        # 保存原始格式
                        original_bold = run.bold
                        original_italic = run.italic
                        original_underline = run.underline
                        original_size = run.font.size
                        
                        # 清空原始run
                        run.text = ""
                        
                        # 为每个字符创建新run
                        for char_idx, char in enumerate(text):
                            # 查找支持该字符的字体
                            font_name = self.font_manager.get_font_for_char(char)
                            
                            new_run = paragraph.add_run(char)
                            
                            if font_name:
                                new_run.font.name = font_name
                                new_run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                                stats['chars_with_font'] += 1
                                stats['used_fonts'].add(font_name)
                            else:
                                stats['chars_without_font'] += 1
                            
                            # 恢复基本格式
                            new_run.bold = original_bold
                            new_run.italic = original_italic
                            new_run.underline = original_underline
                            
                            # 应用随机字符大小
                            if self.enable_random_char_size.get() and original_size:
                                current_size = self._get_random_char_size(original_size.pt, last_char_size, char_size_range)
                                new_run.font.size = Pt(current_size)
                                last_char_size = current_size
                                stats['chars_with_random_size'] += 1
                            elif original_size:
                                # 保持原始大小
                                new_run.font.size = original_size
                                last_char_size = original_size.pt
                            
                            # 应用手写倾斜效果
                            if self.enable_handwriting_effect.get():
                                tilt_angle = simulator.get_char_tilt(char)
                                # 根据强度调整倾斜幅度
                                adjusted_tilt = tilt_angle * max_tilt_multiplier
                                self._apply_character_tilt(new_run, adjusted_tilt)
                            
                            # 应用字符高度位置随机化（限制相邻字符高度落差）
                            char_position = self._get_random_char_position()
                            self._apply_character_position(new_run, char_position)
                            
                            stats['total_chars'] += 1
                
                # 记录趋势数量（用于统计）
                stats['handwriting_trends'] += simulator.char_count_since_correction
            
            # 处理表格
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        # 为表格中的每个单元格段落创建独立模拟器
                        cell_key = f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}"
                        
                        for para_idx, paragraph in enumerate(cell.paragraphs):
                            if f"{cell_key}_para_{para_idx}" not in paragraph_simulators:
                                paragraph_simulators[f"{cell_key}_para_{para_idx}"] = HandwritingSimulator()
                            
                            simulator = paragraph_simulators[f"{cell_key}_para_{para_idx}"]
                            
                            # 应用随机行间距（表格中的段落）
                            if self.enable_random_line_spacing.get() and paragraph.text.strip():
                                random_spacing = random.uniform(line_spacing_min, line_spacing_max)
                                paragraph.paragraph_format.line_spacing = random_spacing
                                stats['lines_with_random_spacing'] += 1
                            
                            # 应用随机行首缩进（表格中的段落）
                            if self.enable_random_indent.get() and paragraph.text.strip():
                                # 在缩进范围内随机选择空格数量
                                indent_spaces = random.randint(indent_min, indent_max)
                                # 在段落开头添加空格
                                if paragraph.runs:
                                    # 如果段落已有内容，在第一个run前插入空格
                                    first_run = paragraph.runs[0]
                                    spaces = " " * indent_spaces
                                    first_run.text = spaces + first_run.text
                                    stats['lines_with_random_indent'] += 1
                                else:
                                    # 如果段落没有内容，添加一个包含空格的run
                                    paragraph.add_run(" " * indent_spaces)
                                    stats['lines_with_random_indent'] += 1
                            
                            # 初始化字符大小跟踪和高度位置跟踪
                            last_char_size = None
                            self.last_char_position = None
                            
                            runs = list(paragraph.runs)
                            for run in runs:
                                text = run.text
                                if text.strip():
                                    # 保存原始格式
                                    original_bold = run.bold
                                    original_italic = run.italic
                                    original_underline = run.underline
                                    original_size = run.font.size
                                    
                                    # 清空原始run
                                    run.text = ""
                                    
                                    # 为每个字符创建新run
                                    for char in text:
                                        font_name = self.font_manager.get_font_for_char(char)
                                        
                                        new_run = paragraph.add_run(char)
                                        
                                        if font_name:
                                            new_run.font.name = font_name
                                            new_run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                                            stats['chars_with_font'] += 1
                                            stats['used_fonts'].add(font_name)
                                        else:
                                            stats['chars_without_font'] += 1
                                        
                                        # 恢复格式
                                        new_run.bold = original_bold
                                        new_run.italic = original_italic
                                        new_run.underline = original_underline
                                        
                                        # 应用随机字符大小
                                        if self.enable_random_char_size.get() and original_size:
                                            current_size = self._get_random_char_size(original_size.pt, last_char_size, char_size_range)
                                            new_run.font.size = Pt(current_size)
                                            last_char_size = current_size
                                            stats['chars_with_random_size'] += 1
                                        elif original_size:
                                            # 保持原始大小
                                            new_run.font.size = original_size
                                            last_char_size = original_size.pt
                                        
                                        # 应用手写倾斜效果
                                        if self.enable_handwriting_effect.get():
                                            tilt_angle = simulator.get_char_tilt(char)
                                            adjusted_tilt = tilt_angle * max_tilt_multiplier
                                            self._apply_character_tilt(new_run, adjusted_tilt)
                                        
                                        # 应用字符高度位置随机化（限制相邻字符高度落差）
                                        char_position = self._get_random_char_position()
                                        self._apply_character_position(new_run, char_position)
                                        
                                        stats['total_chars'] += 1
            
            # 保存文档
            doc.save(output_path)
            
            # 更新UI（在主线程中）
            self.root.after(0, self.conversion_completed, output_path, stats)
            
        except Exception as e:
            # 错误处理（在主线程中）
            self.root.after(0, self.conversion_failed, str(e))
            print(traceback.format_exc())
    
    def _get_random_char_size(self, base_size, last_char_size=None, size_range=0.8):
        """
        获取随机字符大小
        在原有字号加减指定范围的区域随机，且相邻两个字符的字号差距不超过0.5
        """
        if last_char_size is None:
            # 第一个字符，在基础大小±size_range范围内随机
            min_size = max(6, base_size - size_range)  # 最小6pt
            max_size = base_size + size_range
            return random.uniform(min_size, max_size)
        else:
            # 后续字符，确保与上一个字符的差距不超过0.5
            min_size = max(6, last_char_size - 0.5, base_size - size_range)
            max_size = min(last_char_size + 0.5, base_size + size_range)
            return random.uniform(min_size, max_size)
    
    def _get_random_char_position(self):
        """
        获取随机字符高度位置
        限制相邻字符的高度落差在合理范围内（-2到2磅之间，相邻字符差距不超过1磅）
        """
        if self.last_char_position is None:
            # 第一个字符，随机位置
            position = random.uniform(-2.5, 2.5)
        else:
            # 后续字符，确保与上一个字符的高度差距不超过1磅
            min_position = max(-2.5, self.last_char_position - 0.5)
            max_position = min(2.5, self.last_char_position + 0.5)
            position = random.uniform(min_position, max_position)
        
        self.last_char_position = position
        return position
    
    def _apply_character_position(self, run, position):
        """为单个字符应用高度位置效果"""
        try:
            if abs(position) < 0.1:  # 忽略非常小的位置偏移
                return
                
            rpr = run._element.rPr
            if rpr is None:
                rpr = OxmlElement('w:rPr')
                run._element.append(rpr)
            
            # 使用位置偏移设置字符高度
            # 转换为半磅单位（Word中的位置单位）
            position_value = int(position * 2)
            
            position_elem = OxmlElement('w:position')
            position_elem.set(qn('w:val'), str(position_value))
            rpr.append(position_elem)
            
        except Exception as e:
            # 忽略位置应用错误，不影响主要功能
            pass
    
    def _apply_character_tilt(self, run, tilt_angle):
        """为单个字符应用倾斜效果"""
        try:
            if abs(tilt_angle) < 0.1:  # 忽略非常小的倾斜
                return
                
            rpr = run._element.rPr
            if rpr is None:
                rpr = OxmlElement('w:rPr')
                run._element.append(rpr)
            
            # 使用位置偏移模拟倾斜效果
            # 转换为半磅单位（Word中的位置单位）
            position_value = int(tilt_angle * 2)
            
            position = OxmlElement('w:position')
            position.set(qn('w:val'), str(position_value))
            rpr.append(position)
            
        except Exception as e:
            # 忽略倾斜应用错误，不影响主要功能
            pass
    
    def conversion_completed(self, output_path, stats):
        """转换完成"""
        self.is_processing = False
        self.convert_btn.config(state="normal")
        self.update_status("转换完成")
        
        self.log(f"字符级字体替换完成！")
        self.log(f"输出文件: {output_path}")
        self.log(f"总共处理了 {stats['total_chars']} 个字符")
        self.log(f"成功应用字体的字符: {stats['chars_with_font']} 个")
        self.log(f"未找到合适字体的字符: {stats['chars_without_font']} 个")
        self.log(f"使用了 {len(stats['used_fonts'])} 种不同的字体")
        
        if self.enable_handwriting_effect.get():
            self.log(f"手写模拟: 应用了智能倾斜和纠正效果")
        
        if self.enable_random_line_spacing.get():
            self.log(f"随机行间距: 应用了 {stats['lines_with_random_spacing']} 行")
        
        if self.enable_random_char_size.get():
            self.log(f"随机字符大小: 应用了 {stats['chars_with_random_size']} 个字符")
        
        if self.enable_random_indent.get():
            self.log(f"随机行首缩进: 应用了 {stats['lines_with_random_indent']} 行")
        
        if stats['used_fonts']:
            self.log("使用的字体: " + ", ".join(list(stats['used_fonts'])[:5]) + 
                    ("..." if len(stats['used_fonts']) > 5 else ""))
        
        message_text = (
            f"字符级字体替换完成！\n\n"
            f"输出文件: {os.path.basename(output_path)}\n"
            f"总共处理了 {stats['total_chars']} 个字符\n"
            f"成功应用字体的字符: {stats['chars_with_font']} 个\n"
            f"未找到合适字体的字符: {stats['chars_without_font']} 个\n"
            f"使用了 {len(stats['used_fonts'])} 种不同的字体"
        )
        
        if self.enable_handwriting_effect.get():
            message_text += f"\n手写模拟: 已启用智能倾斜效果"
        
        if self.enable_random_line_spacing.get():
            message_text += f"\n随机行间距: 已启用 (力度: {self.line_spacing_strength.get()})"
        
        if self.enable_random_char_size.get():
            message_text += f"\n随机字符大小: 已启用 (力度: {self.char_size_strength.get()})"
        
        if self.enable_random_indent.get():
            message_text += f"\n随机行首缩进: 已启用 (力度: {self.indent_strength.get()})"
        
        messagebox.showinfo("完成", message_text)
    
    def conversion_failed(self, error_msg):
        """转换失败"""
        self.is_processing = False
        self.convert_btn.config(state="normal")
        self.update_status("转换失败")
        
        self.log(f"转换失败: {error_msg}")
        messagebox.showerror("错误", f"转换失败:\n{error_msg}")
    
    def clear_all(self):
        """清空所有输入"""
        self.input_entry.delete(0, tk.END)
        self.output_entry.delete(0, tk.END)
        self.input_file = ""
        self.output_file = ""
        self.log_text.delete(1.0, tk.END)
        self.log("已清空所有输入")
        self.update_status("就绪")
    
    def update_status(self, message):
        """更新状态栏"""
        self.status_var.set(message)
    
    def log(self, message):
        """添加日志"""
        self.log_text.insert(tk.END, f"{message}\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()

def main():
    # 创建主窗口
    root = tk.Tk()
    
    # 创建应用
    app = FontRandomizerApp(root)
    
    # 启动主循环
    root.mainloop()

if __name__ == "__main__":
    main()